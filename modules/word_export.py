from __future__ import annotations

import ast
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from modules.protocol_model import SCHRIFTWESEN_ITEMS

BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "D9D9D9"
MEDIUM_GRAY = "A6A6A6"


def create_word_document(protocol: Dict[str, Any]) -> bytes:
    doc = Document()

    _set_document_defaults(doc)
    _set_page_margins(doc)
    _add_header(doc)

    stammdaten = protocol.get("stammdaten", {})
    buv_nummer = clean_text(stammdaten.get("buv_nummer", ""))

    if buv_nummer:
        title_text = f"{buv_nummer}. Besondere Unterrichtsvorbereitung"
    else:
        title_text = "Besondere Unterrichtsvorbereitung"

    _add_title(doc, title_text)

    _add_stammdaten_block(doc, protocol)

    _add_section_heading(doc, "Einzel-BUV")
    _add_section_einzel_buv(doc, protocol.get("einzel_buv", {}))

    _add_section_heading(doc, "Doppel-BUV")
    _add_section_doppel_buv(doc, protocol.get("doppel_buv", {}))

    _add_section_heading(doc, "Erzieherische Kompetenz")
    _add_section_kompetenzen(doc, protocol)

    _add_section_heading(doc, "Zusammenfassung zur Weiterarbeit")
    _add_summary_block(
        doc,
        "Besprechung der Einzel-BUV",
        protocol.get("einzel_buv", {}).get("zusammenfassung_weiterarbeit", ""),
    )
    _add_summary_block(
        doc,
        "Besprechung der Doppel-BUV",
        protocol.get("doppel_buv", {}).get("zusammenfassung_weiterarbeit", ""),
    )

    _add_section_heading(doc, "Zielvereinbarungen des LAA")
    _add_zielvereinbarungen(doc, protocol)

    _add_signature_lines(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Grundlayout
# ---------------------------------------------------------------------------


def _set_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK

    for style_name in doc.styles:
        try:
            style = doc.styles[style_name]
            if hasattr(style, "font"):
                style.font.name = "Arial"
                style.font.color.rgb = BLACK
        except Exception:
            pass


def _set_page_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)


def _add_header(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header

    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    _clear_cell(left_cell)
    _clear_cell(right_cell)

    left_cell.width = Inches(4.7)
    right_cell.width = Inches(1.8)

    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run("Seminar 45.2 München Stadt\n")
    _format_run(run, size=9, bold=False)

    run = p.add_run("Seminarleitung: Marcus Müller, Seminarrektor")
    _format_run(run, size=9, bold=False)

    logo_path = _find_logo_path()
    if logo_path:
        p_logo = right_cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            run_logo = p_logo.add_run()
            run_logo.add_picture(str(logo_path), width=Cm(2.3))
        except Exception:
            pass

    _remove_table_borders(table)


def _find_logo_path() -> Path | None:
    candidates = [
        Path("Seminar 45.2.png"),
        Path("seminar_logo.png"),
        Path("assets") / "Seminar 45.2.png",
        Path("assets") / "seminar_logo.png",
        Path(__file__).resolve().parent.parent / "Seminar 45.2.png",
        Path(__file__).resolve().parent.parent / "seminar_logo.png",
        Path(__file__).resolve().parent.parent / "assets" / "Seminar 45.2.png",
        Path(__file__).resolve().parent.parent / "assets" / "seminar_logo.png",
    ]

    for candidate in candidates:
        if candidate.exists():
            return candidate

    return None


def _find_signature_path() -> Path | None:
    candidates = [
        Path("Digitale Unterschrift.png"),
        Path("digitale_unterschrift.png"),
        Path("unterschrift.png"),
        Path("Unterschrift.png"),
        Path("signature.png"),
        Path("assets") / "Digitale Unterschrift.png",
        Path("assets") / "digitale_unterschrift.png",
        Path("assets") / "unterschrift.png",
        Path("assets") / "Unterschrift.png",
        Path("assets") / "signature.png",
        Path(__file__).resolve().parent.parent / "Digitale Unterschrift.png",
        Path(__file__).resolve().parent.parent / "digitale_unterschrift.png",
        Path(__file__).resolve().parent.parent / "unterschrift.png",
        Path(__file__).resolve().parent.parent / "Unterschrift.png",
        Path(__file__).resolve().parent.parent / "signature.png",
        Path(__file__).resolve().parent.parent / "assets" / "Digitale Unterschrift.png",
        Path(__file__).resolve().parent.parent / "assets" / "digitale_unterschrift.png",
        Path(__file__).resolve().parent.parent / "assets" / "unterschrift.png",
        Path(__file__).resolve().parent.parent / "assets" / "Unterschrift.png",
        Path(__file__).resolve().parent.parent / "assets" / "signature.png",
    ]

    for candidate in candidates:
        if candidate.exists():
            return candidate

    return None


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(14)

    run = p.add_run(clean_text(text))
    _format_run(run, size=18, bold=True)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)

    run = p.add_run(clean_text(text))
    _format_run(run, size=13, bold=True)

    _add_bottom_border(p, color=MEDIUM_GRAY, size="8")


def _add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(3)

    run = p.add_run(clean_text(text))
    _format_run(run, size=11, bold=True)


def _format_run(run, size: int = 10, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


# ---------------------------------------------------------------------------
# Inhalt
# ---------------------------------------------------------------------------


def _add_stammdaten_block(doc: Document, protocol: Dict[str, Any]) -> None:
    stammdaten = protocol.get("stammdaten", {})
    einzel = protocol.get("einzel_buv", {})
    doppel = protocol.get("doppel_buv", {})
    stunde_1 = doppel.get("stunde_1", {})
    stunde_2 = doppel.get("stunde_2", {})

    _add_section_heading(doc, "Stammdaten")

    rows = [
        ("Name LAA", stammdaten.get("laa_name", "")),
        ("Seminarjahr", stammdaten.get("seminarjahr", "")),
        ("Schule", stammdaten.get("schule", "")),
        ("Datum Einzel-BUV", einzel.get("datum", "")),
        (
            "Fach / Klasse Einzel-BUV",
            _join_nonempty([einzel.get("fach", ""), einzel.get("klasse", "")], " / "),
        ),
        ("Thema Einzel-BUV", einzel.get("thema", "")),
        ("Datum Doppel-BUV", doppel.get("datum", "")),
        (
            "Fach / Klasse Doppel-BUV, 1. Stunde",
            _join_nonempty([stunde_1.get("fach", ""), stunde_1.get("klasse", "")], " / "),
        ),
        ("Thema Doppel-BUV, 1. Stunde", stunde_1.get("thema", "")),
        (
            "Fach / Klasse Doppel-BUV, 2. Stunde",
            _join_nonempty([stunde_2.get("fach", ""), stunde_2.get("klasse", "")], " / "),
        ),
        ("Thema Doppel-BUV, 2. Stunde", stunde_2.get("thema", "")),
    ]

    for label, value in rows:
        _add_label_value_line(doc, label, value)

    bemerkungen = clean_text(stammdaten.get("bemerkungen", ""))
    if bemerkungen:
        _add_label_value_line(doc, "Bemerkungen", bemerkungen)


def _add_section_einzel_buv(doc: Document, einzel: Dict[str, Any]) -> None:
    _add_subheading(doc, "Kurzanalyse des Entwurfs")
    _add_bullet_list(doc, einzel.get("entwurf_analyse", []))

    _add_subheading(doc, "Unterrichtskompetenz – Beobachtungen")
    _add_observations(doc, einzel.get("beobachtungen", {}))


def _add_section_doppel_buv(doc: Document, doppel: Dict[str, Any]) -> None:
    for key, label in [
        ("stunde_1", "Doppel-BUV – 1. Stunde"),
        ("stunde_2", "Doppel-BUV – 2. Stunde"),
    ]:
        stunde = doppel.get(key, {})

        _add_subheading(doc, label)

        _add_label_value_line(
            doc,
            "Fach / Klasse",
            _join_nonempty([stunde.get("fach", ""), stunde.get("klasse", "")], " / "),
        )
        _add_label_value_line(doc, "Thema", stunde.get("thema", ""))

        _add_subheading(doc, f"Kurzanalyse des Entwurfs – {label}")
        _add_bullet_list(doc, stunde.get("entwurf_analyse", []))

        _add_subheading(doc, f"Unterrichtskompetenz – Beobachtungen – {label}")
        _add_observations(doc, stunde.get("beobachtungen", {}))


def _add_section_kompetenzen(doc: Document, protocol: Dict[str, Any]) -> None:
    kompetenzen = protocol.get("kompetenzen", {})
    erzieh = kompetenzen.get("erzieherische_kompetenz", {})

    _add_criteria_box(
        doc,
        "Erzieherische Kompetenz",
        [
            ("Positive Feststellungen", erzieh.get("positive_feststellungen", "")),
            ("Beratungspunkte", erzieh.get("beratungspunkte", "")),
        ],
    )

    _add_section_heading(doc, "Handlungs- und Sachkompetenz")
    _add_schriftwesen_table(doc, protocol)

    _add_text_section_if_present(
        doc,
        "Weitere Hinweise zur Handlungs- und Sachkompetenz",
        kompetenzen.get("handlungs_und_sachkompetenz", ""),
    )

    _add_text_section_if_present(
        doc,
        "Einbringen in Schule und Seminar",
        kompetenzen.get("einbringen_schule_und_seminar", ""),
    )


def _add_observations(doc: Document, observations: Dict[str, Any]) -> None:
    if not observations:
        _add_empty_line(doc)
        return

    added_any = False

    for kriterium, raw_fields in observations.items():
        fields = _normalize_observation_fields(raw_fields)

        positive = clean_text(fields.get("positive_feststellungen", ""))
        beratung = clean_text(fields.get("beratungspunkte", ""))

        # Rohnotizen / Memo werden bewusst NICHT in das Word-Dokument übernommen.
        if not positive and not beratung:
            continue

        _add_criteria_box(
            doc,
            clean_text(kriterium),
            [
                ("Positive Feststellungen", positive),
                ("Beratungspunkte", beratung),
            ],
        )
        added_any = True

    if not added_any:
        _add_empty_line(doc)


def _add_summary_block(doc: Document, title: str, text: Any) -> None:
    cleaned = clean_text(text)

    _add_subheading(doc, title)

    if not cleaned:
        _add_empty_line(doc)
        return

    lines = _to_lines(cleaned)

    for line in lines:
        _add_bullet(doc, line)


def _add_zielvereinbarungen(doc: Document, protocol: Dict[str, Any]) -> None:
    einzel_zv = clean_text(protocol.get("einzel_buv", {}).get("zielvereinbarungen_laa", ""))
    doppel_zv = clean_text(protocol.get("doppel_buv", {}).get("zielvereinbarungen_laa", ""))

    p = doc.add_paragraph()
    run = p.add_run(
        "Bitte tragen Sie Ihre Zielvereinbarungen ein und senden Sie das Dokument anschließend zurück."
    )
    _format_run(run, size=10, italic=True)

    _add_subheading(doc, "Zur Einzel-BUV")
    if einzel_zv:
        _add_multiline_text(doc, einzel_zv)
    else:
        _add_blank_bullets(doc, 4)

    _add_subheading(doc, "Zur Doppel-BUV")
    if doppel_zv:
        _add_multiline_text(doc, doppel_zv)
    else:
        _add_blank_bullets(doc, 4)

def _add_schriftwesen_table(doc: Document, protocol: Dict[str, Any]) -> None:
    kompetenzen = protocol.get("kompetenzen", {})
    schriftwesen = kompetenzen.get("schriftwesen", {})
    doppel_datum = clean_text(protocol.get("doppel_buv", {}).get("datum", ""))

    if doppel_datum:
        title = f"Schriftwesen (vorgelegt am: {doppel_datum})"
    else:
        title = "Schriftwesen (vorgelegt am: __________)"

    _add_subheading(doc, title)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = None
    table.autofit = True

    headers = ["", "OK", "fehlt", "Bemerkung"]

    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        _set_cell_shading(cell, "F2F2F2")
        _set_cell_border(cell, color=LIGHT_GRAY)
        cell.text = header

        for p in cell.paragraphs:
            for run in p.runs:
                _format_run(run, size=10, bold=True)

    for item in SCHRIFTWESEN_ITEMS:
        row_data = schriftwesen.get(item, {})
        status = clean_text(row_data.get("status", "OK"))
        bemerkung = clean_text(row_data.get("bemerkung", ""))

        cells = table.add_row().cells

        cells[0].text = clean_text(item)

        if status == "OK":
            cells[1].text = "x"
            cells[2].text = ""
        elif status == "fehlt":
            cells[1].text = ""
            cells[2].text = "x"
        elif status == "--":
            cells[1].text = "--"
            cells[2].text = ""
        else:
            cells[1].text = ""
            cells[2].text = ""

        cells[3].text = bemerkung

        for cell in cells:
            _set_cell_border(cell, color=LIGHT_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    _format_run(run, size=9.5)

    # Spalten etwas ausrichten
    for row in table.rows:
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

# ---------------------------------------------------------------------------
# Darstellungshelfer
# ---------------------------------------------------------------------------


def _add_label_value_line(doc: Document, label: str, value: Any) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)

    run_label = p.add_run(f"{clean_text(label)}: ")
    _format_run(run_label, size=10.5, bold=True)

    run_value = p.add_run(clean_text(value))
    _format_run(run_value, size=10.5)


def _add_criteria_box(doc: Document, title: str, rows: Iterable[tuple[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_shading(cell, "F7F7F7")
    _set_cell_border(cell, color=LIGHT_GRAY)

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(clean_text(title))
    _format_run(run_title, size=10.8, bold=True)

    for label, value in rows:
        cleaned = clean_text(value)
        if not cleaned:
            continue

        p_label = cell.add_paragraph()
        p_label.paragraph_format.space_before = Pt(2)
        p_label.paragraph_format.space_after = Pt(0)
        run_label = p_label.add_run(f"{clean_text(label)}:")
        _format_run(run_label, size=10, bold=True)

        for line in _to_lines(cleaned):
            p_text = cell.add_paragraph()
            p_text.paragraph_format.left_indent = Cm(0.35)
            p_text.paragraph_format.space_after = Pt(0)
            run_text = p_text.add_run(line)
            _format_run(run_text, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_text_section_if_present(doc: Document, title: str, text: Any) -> None:
    cleaned = clean_text(text)

    if not cleaned:
        return

    _add_subheading(doc, title)
    _add_multiline_text(doc, cleaned)


def _add_multiline_text(doc: Document, text: str) -> None:
    for line in _to_lines(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        _format_run(run, size=10.5)


def _add_bullet_list(doc: Document, points: Any) -> None:
    lines = _to_lines(points)

    if not lines:
        _add_empty_line(doc)
        return

    for line in lines:
        _add_bullet(doc, line)


def _add_bullet(doc: Document, text: Any) -> None:
    cleaned = clean_text(text)

    if not cleaned:
        return

    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.15)
    p.paragraph_format.space_after = Pt(2)

    run_bullet = p.add_run("• ")
    _format_run(run_bullet, size=10.5)

    run_text = p.add_run(cleaned)
    _format_run(run_text, size=10.5)


def _add_empty_line(doc: Document) -> None:
    p = doc.add_paragraph("—")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        _format_run(run, size=10.5)


def _add_blank_bullets(doc: Document, count: int) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("- ")
        _format_run(run, size=10.5)


def _add_signature_lines(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    # Datumszeile
    table.cell(0, 0).text = "München, den ____________________"
    table.cell(0, 1).text = "München, den ____________________"

    # Unterschriftsbild links einfügen
    signature_path = _find_signature_path()

    left_signature_cell = table.cell(1, 0)
    right_signature_cell = table.cell(1, 1)

    _clear_cell(left_signature_cell)
    _clear_cell(right_signature_cell)

    p_sig = left_signature_cell.paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if signature_path:
        try:
            run_sig = p_sig.add_run()
            run_sig.add_picture(str(signature_path), width=Cm(4.0))
        except Exception:
            fallback = p_sig.add_run("________________________")
            _format_run(fallback, size=10)
    else:
        fallback = p_sig.add_run("________________________")
        _format_run(fallback, size=10)

    p_laa = right_signature_cell.paragraphs[0]
    p_laa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_laa = p_laa.add_run("________________________")
    _format_run(run_laa, size=10)

    # Namenszeile
    table.cell(2, 0).text = "Marcus Müller, Seminarrektor"
    table.cell(2, 1).text = "Unterschrift LAA"

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _format_run(run, size=10)


# ---------------------------------------------------------------------------
# Textbereinigung
# ---------------------------------------------------------------------------


def clean_text(value: Any) -> str:
    if value is None:
        return ""

    if isinstance(value, dict):
        value = _dict_to_readable_text(value)

    if isinstance(value, list):
        value = "\n".join(clean_text(item) for item in value if clean_text(item))

    text = str(value)

    # Falls ein Dictionary versehentlich als String im Word-Export landet.
    stripped = text.strip()
    if stripped.startswith("{") and stripped.endswith("}"):
        try:
            parsed = ast.literal_eval(stripped)
            if isinstance(parsed, dict):
                text = _dict_to_readable_text(parsed)
        except Exception:
            pass

    # Markdown-Bold/Italic entfernen.
    text = text.replace("**", "")
    text = text.replace("__", "")

    # Einzelne störende Markdown- oder Listenzeichen am Zeilenanfang bereinigen.
    # Wichtig: Datumsangaben wie 12.05.2026 dürfen nicht als nummerierte Liste behandelt werden.
    lines = []
    for line in text.splitlines():
        line = line.strip()
        line = re.sub(r"^[\-•]\s*", "", line)

        # Nummerierte Listen wie "1. Text" oder "1) Text" bereinigen,
        # aber Datumsangaben wie "12.05.2026" unverändert lassen.
        if not re.match(r"^\d{1,2}\.\d{1,2}\.\d{2,4}$", line):
            line = re.sub(r"^\d+[\.\)]\s+", "", line)

        line = line.strip()

        # Umschließende einfache oder doppelte Anführungszeichen entfernen.
        line = line.strip("\"'")
        lines.append(line)

    text = "\n".join(lines)

    replacements = {
        "positive_feststellungen": "Positive Feststellungen",
        "beratungspunkte": "Beratungspunkte",
        "memo": "Rohnotizen",
        "entwurf_analyse": "Kurzanalyse des Entwurfs",
        "zusammenfassung_weiterarbeit": "Zusammenfassung zur Weiterarbeit",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"[ \t]+", " ", text)

    return text.strip()


def _normalize_observation_fields(raw_fields: Any) -> Dict[str, str]:
    """Sorgt dafür, dass Beobachtungsfelder immer als sauberes Dictionary vorliegen.

    Manchmal kann durch KI-Ausgaben oder ältere JSON-Stände ein Feld nicht als dict,
    sondern als Text/String gespeichert sein. Diese Funktion verhindert technische
    Ausgaben wie {'positive_feststellungen': ...} im Word-Dokument.
    """
    if raw_fields is None:
        return {
            "positive_feststellungen": "",
            "beratungspunkte": "",
            "memo": "",
        }

    if isinstance(raw_fields, dict):
        return {
            "positive_feststellungen": clean_text(raw_fields.get("positive_feststellungen", "")),
            "beratungspunkte": clean_text(raw_fields.get("beratungspunkte", "")),
            "memo": clean_text(raw_fields.get("memo", "")),
        }

    if isinstance(raw_fields, str):
        stripped = raw_fields.strip()

        if stripped.startswith("{") and stripped.endswith("}"):
            try:
                parsed = ast.literal_eval(stripped)
                if isinstance(parsed, dict):
                    return {
                        "positive_feststellungen": clean_text(parsed.get("positive_feststellungen", "")),
                        "beratungspunkte": clean_text(parsed.get("beratungspunkte", "")),
                        "memo": clean_text(parsed.get("memo", "")),
                    }
            except Exception:
                pass

        return {
            "positive_feststellungen": "",
            "beratungspunkte": clean_text(raw_fields),
            "memo": "",
        }

    return {
        "positive_feststellungen": "",
        "beratungspunkte": clean_text(raw_fields),
        "memo": "",
    }


def _dict_to_readable_text(data: Dict[str, Any]) -> str:
    parts: List[str] = []

    positive = data.get("positive_feststellungen")
    beratung = data.get("beratungspunkte")

    if positive:
        parts.append(f"Positive Feststellungen:\n{clean_text(positive)}")

    if beratung:
        parts.append(f"Beratungspunkte:\n{clean_text(beratung)}")

    for key, value in data.items():
        if key in {"positive_feststellungen", "beratungspunkte", "memo"}:
            continue

        cleaned_value = clean_text(value)
        if cleaned_value:
            readable_key = clean_text(key)
            parts.append(f"{readable_key}:\n{cleaned_value}")

    return "\n".join(parts)


def _to_lines(value: Any) -> List[str]:
    cleaned = clean_text(value)

    if not cleaned:
        return []

    lines = []

    for line in cleaned.splitlines():
        line = clean_text(line)
        if line:
            lines.append(line)

    return lines


def _join_nonempty(values: Iterable[Any], separator: str = " / ") -> str:
    cleaned = [clean_text(value) for value in values if clean_text(value)]
    return separator.join(cleaned)


# ---------------------------------------------------------------------------
# Word-XML-Helfer
# ---------------------------------------------------------------------------


def _add_bottom_border(paragraph, color: str = MEDIUM_GRAY, size: str = "6") -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def _set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))

    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)

    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = LIGHT_GRAY) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "nil")


def _clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.clear()
